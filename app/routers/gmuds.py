"""
Rotas para gerenciamento de GMUDs
"""
from fastapi import APIRouter, Depends, HTTPException, status
from starlette.responses import FileResponse, StreamingResponse
from sqlalchemy.orm import Session
from typing import List
import logging
from datetime import datetime
import io

from app.database import get_db
from app.models import GMUD
from app.schemas import GMUDCreate, GMUDUpdate, GMUDResponse
from app.services import GLPIService

logger = logging.getLogger(__name__)
router = APIRouter(prefix="/api/gmuds", tags=["gmuds"])

glpi_service = GLPIService()


@router.post("/", response_model=GMUDResponse, status_code=status.HTTP_201_CREATED)
def criar_gmud(gmud: GMUDCreate, db: Session = Depends(get_db)):
    """
    Cria uma nova GMUD (Gestão de Manutenção)
    
    - **data**: Data da manutenção (YYYY-MM-DD)
    - **equipamento**: Nome do equipamento
    - **descricao**: Descrição da manutenção
    - **justificativa**: Justificativa da manutenção
    - **risco**: Nível de risco (BAIXO, MÉDIO, ALTO)
    """
    # Verificar duplicação
    existe = db.query(GMUD).filter(
        GMUD.data == gmud.data,
        GMUD.equipamento == gmud.equipamento
    ).first()
    
    if existe:
        raise HTTPException(
            status_code=status.HTTP_400_BAD_REQUEST,
            detail="Já existe uma GMUD agendada para este equipamento nesta data"
        )
    
    # Criar chamado no GLPI se configurado
    glpi_result = glpi_service.criar_chamado(gmud.equipamento, gmud.data, gmud.descricao)
    glpi_id = glpi_result.get("glpi_id") if glpi_result else None
    
    # Salvar no banco
    nova_gmud = GMUD(
        data=gmud.data,
        equipamento=gmud.equipamento,
        descricao=gmud.descricao,
        justificativa=gmud.justificativa,
        risco=gmud.risco,
        glpi_id=glpi_id
    )
    
    db.add(nova_gmud)
    db.commit()
    db.refresh(nova_gmud)
    
    logger.info(f"GMUD criada: {nova_gmud.id}")
    return nova_gmud


@router.get("/", response_model=List[GMUDResponse])
def listar_gmuds(
    equipamento: str = None,
    status: str = None,
    db: Session = Depends(get_db)
):
    """
    Lista todas as GMUDs com filtros opcionais
    
    - **equipamento**: Filtrar por equipamento
    - **status**: Filtrar por status (AGENDADO, EM_PROGRESSO, CONCLUIDO, CANCELADO)
    """
    query = db.query(GMUD)
    
    if equipamento:
        query = query.filter(GMUD.equipamento.ilike(f"%{equipamento}%"))
    
    if status:
        query = query.filter(GMUD.status == status)
    
    return query.order_by(GMUD.data).all()


@router.get("/{gmud_id}/documento", tags=["documentos"])
def baixar_documento_gmud(gmud_id: int, db: Session = Depends(get_db)):
    """
    Gera e faz download do documento GMUD em formato Word (.docx)
    """
    try:
        from docx import Document
        from docx.shared import Pt, Inches, RGBColor
        
        # Obter GMUD do banco
        gmud = db.query(GMUD).filter(GMUD.id == gmud_id).first()
        
        if not gmud:
            raise HTTPException(
                status_code=status.HTTP_404_NOT_FOUND,
                detail="GMUD não encontrada"
            )
        
        # Criar documento Word - Modelo profissional REDE AMAZÔNICA (UMA PÁGINA)
        doc = Document()
        
        # Definir margens estreitas para caber em uma página
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.3)
            section.bottom_margin = Inches(0.3)
            section.left_margin = Inches(0.3)
            section.right_margin = Inches(0.3)
        
        # ====== HEADER COM LOGO E DADOS ======
        header_table = doc.add_table(rows=1, cols=2)
        header_table.autofit = False
        
        cell_logo = header_table.rows[0].cells[0]
        cell_logo.text = 'REDE AMAZÔNICA'
        cell_logo.paragraphs[0].runs[0].font.bold = True
        cell_logo.paragraphs[0].runs[0].font.size = Pt(12)
        cell_logo.paragraphs[0].runs[0].font.color.rgb = RGBColor(0, 176, 176)
        cell_logo.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell_logo.paragraphs[0].paragraph_format.space_after = Pt(0)
        
        cell_right = header_table.rows[0].cells[1]
        cell_right.text = 'Código: ________\nRevisão: __  |  Data: __/__/____'
        cell_right.paragraphs[0].alignment = 2
        cell_right.paragraphs[0].runs[0].font.size = Pt(8)
        cell_right.paragraphs[0].paragraph_format.space_before = Pt(0)
        cell_right.paragraphs[0].paragraph_format.space_after = Pt(0)
        
        # Espaço mínimo
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        
        # ====== DADOS DO EQUIPAMENTO ======
        doc.add_heading('DADOS DO EQUIPAMENTO', level=3).runs[0].font.size = Pt(9)
        doc.paragraphs[-1].paragraph_format.space_before = Pt(0)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(1)
        
        table1 = doc.add_table(rows=3, cols=2)
        table1.style = 'Light Grid Accent 1'
        
        table1.rows[0].cells[0].text = 'Serviço'
        table1.rows[0].cells[1].text = 'Data Visitação'
        table1.rows[1].cells[0].text = gmud.equipamento or '___________'
        table1.rows[1].cells[1].text = str(gmud.data) if gmud.data else '___________'
        table1.rows[2].cells[0].text = 'Áreas Envolvidas'
        table1.rows[2].cells[1].text = (gmud.descricao or '______________')[:35]
        
        for row in table1.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(7)
                    paragraph.paragraph_format.space_before = Pt(1)
                    paragraph.paragraph_format.space_after = Pt(1)
                    paragraph.paragraph_format.line_spacing = 1.0
        
        # ====== DADOS DA MANUTENÇÃO ======
        doc.add_heading('DADOS DA MANUTENÇÃO', level=3).runs[0].font.size = Pt(9)
        doc.paragraphs[-1].paragraph_format.space_before = Pt(2)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(1)
        
        table2 = doc.add_table(rows=2, cols=2)
        table2.style = 'Light Grid Accent 1'
        
        table2.rows[0].cells[0].text = 'Justificativa'
        table2.rows[0].cells[1].text = (gmud.justificativa or '______________')[:30]
        table2.rows[1].cells[0].text = 'Descrição da Manutenção'
        table2.rows[1].cells[1].text = (gmud.descricao or '______________')[:30]
        
        for row in table2.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(7)
                    paragraph.paragraph_format.space_before = Pt(1)
                    paragraph.paragraph_format.space_after = Pt(1)
                    paragraph.paragraph_format.line_spacing = 1.0
        
        # ====== INFORMAÇÕES ADICIONAIS ======
        doc.add_heading('INFORMAÇÕES ADICIONAIS', level=3).runs[0].font.size = Pt(9)
        doc.paragraphs[-1].paragraph_format.space_before = Pt(2)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(1)
        
        table3 = doc.add_table(rows=3, cols=2)
        table3.style = 'Light Grid Accent 1'
        
        table3.rows[0].cells[0].text = 'Tempo Estimado'
        table3.rows[0].cells[1].text = '___h ___min'
        table3.rows[1].cells[0].text = 'Análise Risco/Impacto'
        table3.rows[1].cells[1].text = gmud.risco or 'BAIXO'
        table3.rows[2].cells[0].text = 'Tipo: ☐ Preventiva  ☐ Emergencial'
        table3.rows[2].cells[1].text = 'Com. Social: ☐ Sim  ☐ Não'
        
        for row in table3.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(7)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 1.0
        
        # ====== ASSINATURAS ======
        doc.add_heading('ASSINATURAS', level=3).runs[0].font.size = Pt(9)
        doc.paragraphs[-1].paragraph_format.space_before = Pt(2)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(1)
        
        sig_table = doc.add_table(rows=2, cols=3)
        sig_table.style = 'Light Grid Accent 1'
        
        sig_table.rows[0].cells[0].text = '___/___/___\nAdm. / Cambista'
        sig_table.rows[0].cells[1].text = '___/___/___\nSupervisor'
        sig_table.rows[0].cells[2].text = '___/___/___\nCambista Turno'
        sig_table.rows[1].cells[0].text = '___/___/___\nSolicitante'
        sig_table.rows[1].cells[1].text = '___/___/___\nTecnólogo'
        sig_table.rows[1].cells[2].text = '___/___/___\nDiretor'
        
        for row in sig_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(6)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 0.9
        
        # ====== RECURSOS UTILIZADOS ======
        doc.add_heading('RECURSOS UTILIZADOS', level=3).runs[0].font.size = Pt(9)
        doc.paragraphs[-1].paragraph_format.space_before = Pt(2)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(1)
        
        recurso_table = doc.add_table(rows=1, cols=2)
        recurso_table.style = 'Light Grid Accent 1'
        recurso_table.rows[0].cells[0].text = 'Descrição'
        recurso_table.rows[0].cells[1].text = '___________________________'
        
        for row in recurso_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(7)
                    paragraph.paragraph_format.space_before = Pt(1)
                    paragraph.paragraph_format.space_after = Pt(1)
        
        # ====== DOCUMENTOS CRIADOS/ATUALIZADOS ======
        doc.add_heading('DOCUMENTOS CRIADOS/ATUALIZADOS', level=3).runs[0].font.size = Pt(9)
        doc.paragraphs[-1].paragraph_format.space_before = Pt(2)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(1)
        
        doc_table = doc.add_table(rows=1, cols=2)
        doc_table.style = 'Light Grid Accent 1'
        doc_table.rows[0].cells[0].text = 'Descrição'
        doc_table.rows[0].cells[1].text = 'Assinatura/Carimbo Validador'
        
        for row in doc_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(7)
                    paragraph.paragraph_format.space_before = Pt(1)
                    paragraph.paragraph_format.space_after = Pt(1)
        
        # ====== ACEITE RESPONSÁVEL ======
        doc.add_heading('ACEITE RESPONSÁVEL', level=3).runs[0].font.size = Pt(9)
        doc.paragraphs[-1].paragraph_format.space_before = Pt(2)
        doc.paragraphs[-1].paragraph_format.space_after = Pt(1)
        
        aceite_table = doc.add_table(rows=1, cols=4)
        aceite_table.style = 'Light Grid Accent 1'
        
        aceite_table.rows[0].cells[0].text = 'Solicitante'
        aceite_table.rows[0].cells[1].text = 'Resp. Setor\n(Eng.)'
        aceite_table.rows[0].cells[2].text = 'Ger.\nEngenharia'
        aceite_table.rows[0].cells[3].text = 'Diretor\nTecnologia'
        
        for row in aceite_table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.size = Pt(6)
                    paragraph.paragraph_format.space_before = Pt(0)
                    paragraph.paragraph_format.space_after = Pt(0)
                    paragraph.paragraph_format.line_spacing = 0.9
        
        # Salvar em memória
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        # Nome do arquivo
        equipamento_limpo = (gmud.equipamento or 'Equipamento').replace(' ', '_').replace('/', '-')
        nome_arquivo = f"GMUD_{gmud.id}_{equipamento_limpo}.docx"
        
        # Retornar usando StreamingResponse para buffer em memória
        return StreamingResponse(
            iter([buffer.getvalue()]),
            media_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            headers={"Content-Disposition": f"attachment; filename={nome_arquivo}"}
        )
    
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Erro ao gerar documento: {str(e)}", exc_info=True)
        raise HTTPException(
            status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
            detail=f"Erro ao gerar documento: {str(e)}"
        )


@router.get("/{gmud_id}", response_model=GMUDResponse)
def obter_gmud(gmud_id: int, db: Session = Depends(get_db)):
    """Obtém detalhes de uma GMUD específica"""
    gmud = db.query(GMUD).filter(GMUD.id == gmud_id).first()
    
    if not gmud:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="GMUD não encontrada"
        )
    
    return gmud


@router.put("/{gmud_id}", response_model=GMUDResponse)
def atualizar_gmud(
    gmud_id: int,
    update: GMUDUpdate,
    db: Session = Depends(get_db)
):
    """Atualiza o status de uma GMUD"""
    gmud = db.query(GMUD).filter(GMUD.id == gmud_id).first()
    
    if not gmud:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="GMUD não encontrada"
        )
    
    gmud.status = update.status
    db.commit()
    db.refresh(gmud)
    
    logger.info(f"GMUD {gmud_id} atualizada para status: {update.status}")
    return gmud


@router.delete("/{gmud_id}", status_code=status.HTTP_204_NO_CONTENT)
def deletar_gmud(gmud_id: int, db: Session = Depends(get_db)):
    """Deleta uma GMUD"""
    gmud = db.query(GMUD).filter(GMUD.id == gmud_id).first()
    
    if not gmud:
        raise HTTPException(
            status_code=status.HTTP_404_NOT_FOUND,
            detail="GMUD não encontrada"
        )
    
    db.delete(gmud)
    db.commit()
    
    logger.info(f"GMUD {gmud_id} deletada")


@router.get("/dashboard/resumo", tags=["dashboard"])
def resumo_dashboard(db: Session = Depends(get_db)):
    """Retorna resumo do dashboard"""
    total = db.query(GMUD).count()
    agendadas = db.query(GMUD).filter(GMUD.status == "AGENDADO").count()
    em_progresso = db.query(GMUD).filter(GMUD.status == "EM_PROGRESSO").count()
    concluidas = db.query(GMUD).filter(GMUD.status == "CONCLUIDO").count()
    canceladas = db.query(GMUD).filter(GMUD.status == "CANCELADO").count()
    
    return {
        "total": total,
        "agendadas": agendadas,
        "em_progresso": em_progresso,
        "concluidas": concluidas,
        "canceladas": canceladas,
        "taxa_conclusao": (concluidas / total * 100) if total > 0 else 0
    }
