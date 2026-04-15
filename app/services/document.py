"""
Serviço de geração de documentos GMUD em formato Word
"""
from datetime import datetime
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class GMUDDocumentService:
    """Serviço responsável por gerar documentos GMUD em Word"""
    
    @staticmethod
    def gerar_documento_gmud(gmud_data: dict) -> BytesIO:
        """
        Gera um documento Word com os dados do GMUD
        
        Args:
            gmud_data: Dicionário com os dados do GMUD
                {
                    "id": int,
                    "data": str (YYYY-MM-DD),
                    "equipamento": str,
                    "descricao": str,
                    "justificativa": str,
                    "risco": str (BAIXO/MÉDIO/ALTO),
                    "status": str,
                    "criado_em": str
                }
        
        Returns:
            BytesIO: Arquivo Word em memória pronto para download
        """
        # Criar documento
        doc = Document()
        
        # Configurar largura da página
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(0.75)
            section.right_margin = Inches(0.75)
        
        # Cabeçalho com logo/título
        title = doc.add_heading('GMUD - GESTÃO DE MANUTENÇÃO', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(31, 78, 121)
        
        subtitle = doc.add_heading('Plano de Manutenção Preventiva', level=2)
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        subtitle_run = subtitle.runs[0]
        subtitle_run.font.color.rgb = RGBColor(79, 129, 189)
        
        # Informações do documento
        info_para = doc.add_paragraph()
        info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        info_run = info_para.add_run(
            f"Documento gerado em: {datetime.now().strftime('%d/%m/%Y às %H:%M:%S')}\n"
            f"GMUD #{gmud_data.get('id', 'N/A')}"
        )
        info_run.font.size = Pt(9)
        info_run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # Espaço
        
        # Tabela com informações principais
        table = doc.add_table(rows=6, cols=2)
        table.style = 'Light Grid Accent 1'
        
        # Configurar largura das colunas
        for row in table.rows:
            row.cells[0].width = Inches(2)
            row.cells[1].width = Inches(4)
        
        # Preenchimento da tabela
        dados_tabela = [
            ('Data da Manutenção', GMUDDocumentService._formatar_data(gmud_data.get('data', 'N/A'))),
            ('Equipamento', gmud_data.get('equipamento', 'N/A')),
            ('Status', GMUDDocumentService._formatar_status(gmud_data.get('status', 'N/A'))),
            ('Nível de Risco', GMUDDocumentService._formatar_risco(gmud_data.get('risco', 'N/A'))),
            ('Data de Criação', GMUDDocumentService._formatar_data_hora(gmud_data.get('criado_em', 'N/A'))),
            ('Descrição Breve', gmud_data.get('descricao', 'N/A')[:100]),
        ]
        
        for i, (chave, valor) in enumerate(dados_tabela):
            cells = table.rows[i].cells
            
            # Estilo da célula de rótulo
            label_para = cells[0].paragraphs[0]
            label_para.text = chave
            label_run = label_para.runs[0]
            label_run.font.bold = True
            label_run.font.color.rgb = RGBColor(31, 78, 121)
            label_run.font.size = Pt(11)
            
            # Estilo da célula de valor
            value_para = cells[1].paragraphs[0]
            value_para.text = str(valor)
            value_run = value_para.runs[0]
            value_run.font.size = Pt(11)
        
        # Seção de descrição
        doc.add_paragraph()
        doc.add_heading('Descrição Detalhada da Manutenção', level=3)
        desc_para = doc.add_paragraph(gmud_data.get('descricao', 'Sem descrição'))
        desc_para_format = desc_para.paragraph_format
        desc_para_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Seção de justificativa
        doc.add_paragraph()
        doc.add_heading('Justificativa da Manutenção', level=3)
        just_para = doc.add_paragraph(gmud_data.get('justificativa', 'Sem justificativa'))
        just_para_format = just_para.paragraph_format
        just_para_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
        
        # Espaço para assinatura
        doc.add_paragraph()
        doc.add_paragraph()
        
        sig_table = doc.add_table(rows=2, cols=2)
        sig_table.autofit = False
        
        for row in sig_table.rows:
            row.cells[0].width = Inches(3)
            row.cells[1].width = Inches(3)
        
        sig_cells = sig_table.rows[0].cells
        sig_cells[0].text = "_" * 40
        sig_cells[1].text = "_" * 40
        
        assinatura = sig_table.rows[1].cells
        assinatura[0].text = "Técnico Responsável"
        assinatura[1].text = "Gestor de Manutenção"
        
        for cell in sig_table.rows[1].cells:
            for para in cell.paragraphs:
                para_run = para.runs[0]
                para_run.font.size = Pt(9)
                para_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Rodapé
        doc.add_paragraph()
        footer_para = doc.add_paragraph(
            "Este documento foi gerado automaticamente pelo Sistema GMUD. "
            "Mantenha este comprovante para fins de auditoria e rastreamento."
        )
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.runs[0]
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        footer_run.italic = True
        
        # Salvar documento em memória
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)
        
        return doc_io
    
    @staticmethod
    def _formatar_data(data_str: str) -> str:
        """Formata data do formato YYYY-MM-DD para DD/MM/YYYY"""
        try:
            if not data_str or data_str == 'N/A':
                return 'N/A'
            data_obj = datetime.strptime(data_str, '%Y-%m-%d')
            return data_obj.strftime('%d/%m/%Y')
        except (ValueError, TypeError):
            return str(data_str)
    
    @staticmethod
    def _formatar_data_hora(data_hora_str: str) -> str:
        """Formata data e hora do ISO format para DD/MM/YYYY HH:MM"""
        try:
            if not data_hora_str or data_hora_str == 'N/A':
                return 'N/A'
            # Remover 'Z' se presente (UTC indicator)
            data_hora_str = data_hora_str.replace('Z', '').split('.')[0]
            data_obj = datetime.fromisoformat(data_hora_str)
            return data_obj.strftime('%d/%m/%Y %H:%M')
        except (ValueError, TypeError, AttributeError):
            return str(data_hora_str)
    
    @staticmethod
    def _formatar_status(status: str) -> str:
        """Formata o status para apresentação visual"""
        status_map = {
            'AGENDADO': '⏰ Agendado',
            'EM_PROGRESSO': '🔄 Em Progresso',
            'CONCLUIDO': '✓ Concluído',
            'CANCELADO': '✗ Cancelado'
        }
        return status_map.get(status, status)
    
    @staticmethod
    def _formatar_risco(risco: str) -> str:
        """Formata o nível de risco para apresentação visual"""
        risco_map = {
            'BAIXO': '🟢 Baixo',
            'MÉDIO': '🟡 Médio',
            'ALTO': '🔴 Alto'
        }
        return risco_map.get(risco, risco)
